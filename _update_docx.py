import os
from docx import Document
src = r'archive\RB_Business_Profile.docx'
out = r'RB_Business_Profile.docx'
doc=Document(src)
old_new=[
 ('127 Fox Street & Eloff Street','36 Houghton Drive, Houghton Estate, Block D'),
 ('127 Fox St & Eloff St','36 Houghton Drive, Houghton Estate, Block D'),
 ('Johannesburg CBD 2000','Johannesburg, Gauteng'),
 ('Johannesburg CBD','Houghton Estate, Johannesburg')
]
changes=0
for p in doc.paragraphs:
    txt=p.text
    nt=txt
    for o,n in old_new:
        if o in nt:
            nt=nt.replace(o,n)
    if nt!=txt:
        if p.runs:
            p.runs[0].text=nt
            for r in p.runs[1:]:
                r.text=''
        else:
            p.text=nt
        changes+=1
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                txt=p.text
                nt=txt
                for o,n in old_new:
                    if o in nt:
                        nt=nt.replace(o,n)
                if nt!=txt:
                    if p.runs:
                        p.runs[0].text=nt
                        for r in p.runs[1:]: r.text=''
                    else:
                        p.text=nt
                    changes+=1
print('changes',changes)
doc.save(out)
print('saved',out,os.path.getsize(out))
